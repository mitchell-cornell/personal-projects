import csv
import docx
from docx.shared import Pt
from docx.shared import Inches

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Avenir'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

more = {False: ', ', True: ', and'}


smoltext = 9
with open('Consent for Name and Organization Appearing on Letter in Support of UMOJA (Responses) - Cleaned.csv', newline='') as names:
    reader = csv.reader(names)
    first = True
    for row in reader:
        if first:
            first = False
            continue
        par = doc.add_paragraph(row[2].strip())
        counter = 3
        while True:
            run = par.add_run(" " + row[counter+1].strip() + ', ' + row[counter].strip())
            run.font.size = Pt(smoltext)
            counter += 2
            if counter > 7 or row[counter] == '':
                break
            
            conj = par.add_run(more[counter > 5 or row[counter+2]== ''])
            conj.font.size = Pt(smoltext)

doc.save("Names.docx")

